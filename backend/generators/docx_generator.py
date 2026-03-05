"""
DOCX resume generator.
Produces a clean, professionally formatted Word document matching the LaTeX template style.
Uses the same structured resume data format as the LaTeX generator.
"""

from io import BytesIO
from typing import Dict
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _apply_bullets(resume_structured: dict, accepted_bullets: Dict[str, str]) -> dict:
    import copy
    result = copy.deepcopy(resume_structured)
    for section in result.get("sections", []):
        for entry in section.get("entries", []):
            for bullet in entry.get("bullets", []):
                bid = bullet["bullet_id"]
                if bid in accepted_bullets and accepted_bullets[bid] != "original":
                    bullet["text"] = accepted_bullets[bid]
    return result


def _get_entry_fields(entry: dict) -> dict:
    """Extract company/role/dates/location with fallback."""
    company = entry.get("company", "")
    role = entry.get("role", "")
    location = entry.get("location", "")
    dates = entry.get("dates", "")

    if not company:
        # Fallback: use header as company
        header = entry.get("header", "")
        lines = [l.strip() for l in header.split("\n") if l.strip()]
        company = lines[0] if lines else header
        role = role or (lines[1] if len(lines) > 1 else "")

    return {"company": company, "role": role, "location": location, "dates": dates}


def _add_section_header(doc: Document, title: str):
    """Add a bold, underlined section header matching the LaTeX titlerule style."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(11)
    # Add bottom border to simulate \titlerule
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_two_column_line(doc: Document, left: str, right: str, left_bold=False, left_italic=False):
    """Add a paragraph with left text and right-aligned text (simulating \hfill)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)

    # Left run
    left_run = p.add_run(left)
    left_run.bold = left_bold
    left_run.italic = left_italic
    left_run.font.size = Pt(10)

    if right:
        # Tab to right margin
        p.add_run("\t")
        right_run = p.add_run(right)
        right_run.font.size = Pt(10)

        # Set tab stop at right margin
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:pos"), "9360")  # ~6.5 inches in twips
        tabs.append(tab)
        pPr.append(tabs)

    return p


def generate_docx(resume_structured: dict, accepted_bullets: Dict[str, str]) -> bytes:
    """Generate a DOCX matching the LaTeX template style."""
    resume = _apply_bullets(resume_structured, accepted_bullets)
    doc = Document()

    # Page margins (matching 0.5in LaTeX margins)
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Default paragraph spacing
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(2)

    # ---- Header ----
    name = resume.get("name", "")
    if name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(name)
        run.bold = True
        run.font.size = Pt(18)

    title_line = resume.get("title", "")
    if title_line:
        p = doc.add_paragraph(title_line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        for r in p.runs:
            r.font.size = Pt(10)

    contact = resume.get("contact", {})
    location = contact.get("location", "")
    if location:
        p = doc.add_paragraph(location)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        for r in p.runs:
            r.font.size = Pt(10)

    contact_parts = [v for v in [
        contact.get("phone"),
        contact.get("email"),
        contact.get("linkedin"),
    ] if v]
    if contact_parts:
        p = doc.add_paragraph("  |  ".join(contact_parts))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        for r in p.runs:
            r.font.size = Pt(10)

    # ---- Summary ----
    summary = resume.get("summary", "")
    if summary:
        _add_section_header(doc, "Professional Summary")
        p = doc.add_paragraph(summary)
        p.paragraph_format.space_after = Pt(4)

    # ---- Sections ----
    section_order = ["experience", "projects", "education", "skills", "certifications", "other"]
    sections_by_type: Dict[str, list] = {}
    for sec in resume.get("sections", []):
        stype = sec.get("type", "other")
        sections_by_type.setdefault(stype, []).append(sec)

    for stype in section_order:
        for sec in sections_by_type.get(stype, []):
            if stype == "experience":
                _add_section_header(doc, "Work Experience")
                for entry in sec.get("entries", []):
                    _add_experience_entry(doc, entry)
            elif stype == "education":
                _add_section_header(doc, "Education")
                for entry in sec.get("entries", []):
                    _add_education_entry(doc, entry)
            elif stype == "skills":
                _add_section_header(doc, "Technical Skills")
                _add_skills_entries(doc, sec.get("entries", []))
            elif stype == "certifications":
                _add_section_header(doc, "Certifications")
                _add_certifications_entries(doc, sec.get("entries", []))
            else:
                _add_section_header(doc, sec.get("title", "Section"))
                for entry in sec.get("entries", []):
                    _add_generic_entry(doc, entry)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_experience_entry(doc: Document, entry: dict):
    fields = _get_entry_fields(entry)
    company = fields["company"]
    role = fields["role"]
    location = fields["location"]
    dates = fields["dates"]

    # Company line (bold) + location right-aligned
    _add_two_column_line(doc, company, location, left_bold=True)

    # Role line (italic) + dates right-aligned
    if role or dates:
        _add_two_column_line(doc, role, dates, left_italic=True)

    # Bullets
    for b in entry.get("bullets", []):
        text = b.get("text", "").strip()
        if text:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Inches(0.2)
            run = p.add_run(text)
            run.font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _add_education_entry(doc: Document, entry: dict):
    fields = _get_entry_fields(entry)
    school = fields["company"]
    degree = fields["role"]
    dates = fields["dates"]

    _add_two_column_line(doc, school, dates, left_bold=True)

    if degree:
        p = doc.add_paragraph(degree)
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.size = Pt(10)

    for b in entry.get("bullets", []):
        text = b.get("text", "").strip()
        if text:
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(1)
            for r in p.runs:
                r.font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _add_skills_entries(doc: Document, entries: list):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    first = True
    for entry in entries:
        for b in entry.get("bullets", []):
            text = b.get("text", "").strip()
            if not text:
                continue
            if not first:
                p.add_run("  ").font.size = Pt(10)
                # New line within paragraph for each skill category
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
            if ":" in text:
                cat, _, items = text.partition(":")
                r = p.add_run(cat.strip() + ": ")
                r.bold = True
                r.font.size = Pt(10)
                r2 = p.add_run(items.strip())
                r2.font.size = Pt(10)
            else:
                r = p.add_run(text)
                r.font.size = Pt(10)
            first = False

        # Handle entries without bullets (header as skill line)
        if not entry.get("bullets"):
            header = entry.get("header", "").strip() or entry.get("company", "").strip()
            if header:
                if not first:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(2)
                if ":" in header:
                    cat, _, items = header.partition(":")
                    r = p.add_run(cat.strip() + ": ")
                    r.bold = True
                    r.font.size = Pt(10)
                    r2 = p.add_run(items.strip())
                    r2.font.size = Pt(10)
                else:
                    r = p.add_run(header)
                    r.font.size = Pt(10)
                first = False


def _add_certifications_entries(doc: Document, entries: list):
    certs = []
    for entry in entries:
        h = entry.get("company", "").strip() or entry.get("header", "").strip()
        if h:
            certs.append(h)
        for b in entry.get("bullets", []):
            text = b.get("text", "").strip()
            if text:
                certs.append(text)
    if certs:
        p = doc.add_paragraph(", ".join(certs))
        p.paragraph_format.space_after = Pt(4)
        for r in p.runs:
            r.font.size = Pt(10)


def _add_generic_entry(doc: Document, entry: dict):
    header = entry.get("company", "") or entry.get("header", "")
    if header:
        p = doc.add_paragraph()
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)

    for b in entry.get("bullets", []):
        text = b.get("text", "").strip()
        if text:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(text)
            run.font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
