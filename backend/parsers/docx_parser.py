from docx import Document
from io import BytesIO


def parse_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes. Returns raw text including table content."""
    doc = Document(BytesIO(file_bytes))
    lines = []

    # Extract paragraph text
    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text)

    # Extract table text (common in resume layouts)
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                lines.append(row_text)

    return "\n".join(lines)
